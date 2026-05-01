from __future__ import annotations

import argparse
import datetime as dt
import shutil
from pathlib import Path

import requests
from docx import Document


def ensure_output_doc(template: Path, output_dir: Path, date: dt.date) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"study_result_{date.isoformat()}.docx"
    if not output_path.exists():
        shutil.copy2(template, output_path)
    return output_path


def append_row(doc_path: Path, row_values: list[str], table_index: int = 0) -> None:
    doc = Document(doc_path)
    if not doc.tables:
        raise ValueError("문서에 표가 없습니다. 표가 포함된 템플릿을 사용하세요.")

    table = doc.tables[table_index]
    if len(row_values) != len(table.columns):
        raise ValueError(
            f"입력 컬럼 수({len(row_values)})와 표 컬럼 수({len(table.columns)})가 다릅니다."
        )

    row = table.add_row()
    for idx, value in enumerate(row_values):
        row.cells[idx].text = value

    doc.save(doc_path)


def make_local_share_link(path: Path) -> str:
    return f"file://{path.resolve()}"


def upload_to_onedrive(
    file_path: Path,
    access_token: str,
    drive_id: str,
    folder_path: str = "study-results",
) -> dict:
    upload_url = (
        f"https://graph.microsoft.com/v1.0/drives/{drive_id}"
        f"/root:/{folder_path}/{file_path.name}:/content"
    )
    with file_path.open("rb") as f:
        response = requests.put(
            upload_url,
            headers={"Authorization": f"Bearer {access_token}"},
            data=f,
            timeout=60,
        )
    response.raise_for_status()
    return response.json()


def create_onedrive_share_link(
    drive_id: str,
    item_id: str,
    access_token: str,
    scope: str = "organization",
    link_type: str = "view",
) -> str:
    create_link_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/createLink"
    body = {"type": link_type, "scope": scope}
    response = requests.post(
        create_link_url,
        headers={
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json",
        },
        json=body,
        timeout=60,
    )
    response.raise_for_status()
    return response.json()["link"]["webUrl"]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="연구회 스터디 결과를 날짜별 DOCX 파일로 저장")
    parser.add_argument("--template", required=True, help="원본 템플릿 docx 경로")
    parser.add_argument("--output-dir", default="outputs", help="출력 폴더")
    parser.add_argument("--date", default=dt.date.today().isoformat(), help="저장 날짜(YYYY-MM-DD)")
    parser.add_argument("--row", required=True, help="콤마(,)로 구분된 한 줄 데이터")
    parser.add_argument("--table-index", type=int, default=0, help="입력할 표 인덱스")

    parser.add_argument("--upload-onedrive", action="store_true", help="OneDrive 업로드 실행")
    parser.add_argument("--access-token", help="Microsoft Graph access token")
    parser.add_argument("--drive-id", help="업로드 대상 drive id")
    parser.add_argument("--folder-path", default="study-results", help="OneDrive 폴더 경로")
    parser.add_argument("--link-scope", default="organization", choices=["anonymous", "organization"], help="공유 링크 scope")
    parser.add_argument("--link-type", default="view", choices=["view", "edit"], help="공유 링크 타입")

    return parser.parse_args()


def main() -> None:
    args = parse_args()

    template = Path(args.template)
    if not template.exists():
        raise FileNotFoundError(f"템플릿 파일을 찾을 수 없습니다: {template}")

    date = dt.date.fromisoformat(args.date)
    output_doc = ensure_output_doc(template, Path(args.output_dir), date)
    row_values = [part.strip() for part in args.row.split(",")]

    append_row(output_doc, row_values, table_index=args.table_index)
    print(f"저장 완료: {output_doc}")
    print(f"다운로드 링크(로컬): {make_local_share_link(output_doc)}")

    if args.upload_onedrive:
        if not args.access_token or not args.drive_id:
            raise ValueError("OneDrive 업로드 시 --access-token, --drive-id가 필요합니다.")

        uploaded = upload_to_onedrive(
            file_path=output_doc,
            access_token=args.access_token,
            drive_id=args.drive_id,
            folder_path=args.folder_path,
        )
        share_url = create_onedrive_share_link(
            drive_id=args.drive_id,
            item_id=uploaded["id"],
            access_token=args.access_token,
            scope=args.link_scope,
            link_type=args.link_type,
        )
        print(f"OneDrive 저장 완료: {uploaded.get('webUrl', '(webUrl 없음)')}")
        print(f"공유 링크: {share_url}")


if __name__ == "__main__":
    main()
